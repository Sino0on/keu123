import time

from django.shortcuts import render, redirect, get_object_or_404, reverse
from .models import *
from .forms import *
from django.views import generic
from django.contrib.auth import authenticate, login
from .filters import *
from django.contrib.auth.decorators import login_required, permission_required
import datetime
from docx import Document
from django.core.files.base import ContentFile, File


def home(request):
    return render(request, 'home.html', {'educations': Education.objects.all()})


@login_required(login_url='login')
def disciplines(request, pk):
    # indexes = Index.objects.all()
    # filter = IndexFilter(request.GET, queryset=indexes)
    # indexes = filter.qs
    disciplines = Discipline.objects.filter(education=pk, teacher=request.user)
    filter = IndexFilter(request.GET, queryset=disciplines)
    try:
        indexes = Index.objects.filter(data=request.GET["indexes__data"], teacher=request.user)
    except:
        indexes = Index.objects.filter(data=datetime.date.today().year, teacher=request.user).distinct()
    disciplines = filter.qs.distinct()
    return render(request, 'disciplines.html', {'disciplines': disciplines, 'form': filter, 'indexes': indexes})


@login_required(login_url='login')
def discipline(request, pk):
    return render(request, 'discipline.html', {'discipline': Discipline.objects.get(id=pk)})


@permission_required(perm="is_superuser", login_url='login')
def discipline_create(request):
    if request.method == "POST":
        form = DisciplineCreateForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save()
            print(instance.id)
            for i in instance.teacher.all():
                try:
                    index = Index(discripline=instance, data=datetime.date.today().year, teacher=i)
                    index.save()
                except:
                    pass
            return redirect('/')
        return render(request, 'discipline_create.html', {'form': DisciplineCreateForm})
    else:
        return render(request, 'discipline_create.html', {'form': DisciplineCreateForm})


def register(request):  # функция регистрации
    print('da')
    if request.method == 'POST':  # Проверка запроса на пост
        form = UserRegisterForm(request.POST, request.FILES)  # присваиваем форму для данных
        print('POST')
        if form.is_valid():  # Проверка на валидность
            form.save()  # Сохранение в базу
            print('VAlid')
            username = form.cleaned_data['username']
            password = form.cleaned_data['password1']
            user = authenticate(username=username, password=password)  # авторизация юзера
            print('hz')
            login(request, user)  # авторизация юзера
            return redirect('/')  # переадресация на главную страничку
        else:
            print(form.errors)
    else:  # если это запрос не пост
        form = UserRegisterForm()  # Присваивание форму

    context = {'form': form}  # контекст для передачи данных для шаблона
    return render(request, 'register.html', context)


def createnewyear(request):
    print('Print')
    for i in User.objects.all():
        print(i)
        for j in i.disciplines.all():
            print('OK', j)
            try:
                form = Index(discripline=j, data=datetime.date.today().year, teacher=i)
                form.save()
            except:
                pass
            print(i, j)
    return redirect('/')


def updateindex(request, pk):
    print('кирди')
    index = get_object_or_404(Index, id=pk)
    print('тапты')
    if index.data <= datetime.date.today().year:
        if request.user.is_superuser:
            form = IndexAdminForm(instance=index)
        else:
            form = IndexForm
        if request.method == "POST":
            print('post')
            if request.user.is_superuser:
                form = IndexAdminForm(request.POST, instance=index)
                if form.is_valid():
                    form.save()
                    return redirect(reverse('disciplines',  args=(index.discripline.education.id, )))
                print(form.errors)
            else:
                print('adam')
                form = IndexForm(request.POST, instance=index)
                if form.is_valid():
                    print('valid')
                    form.save()
                    return redirect(reverse('disciplines',  args=(index.discripline.education.id, )))
                print(form.errors)
        return render(request, 'index_update.html', {'form': form, 'index': index })
    return redirect(reverse('disciplines',  args=(index.discripline.education.id, )))


def user_detail(request, pk):
    account = User.objects.get(id=pk)
    disciplines = Discipline.objects.filter(teacher=account)
    educations = Education.objects.all()
    print(disciplines)
    filter = IndexFilter(request.GET, queryset=disciplines)
    try:
        indexes = Index.objects.filter(data=request.GET["indexes__data"], teacher=account)
    except:
        indexes = Index.objects.filter(data=datetime.date.today().year, teacher=account).distinct()
    disciplines = filter.qs.distinct()
    print(disciplines)
    wordform = WordGenForm()
    return render(request, 'user_detail.html', {'account': account, 'disciplines': disciplines, 'form': filter, 'indexes': indexes, 'educations': educations, 'wordform': wordform})


def user_list(request):
    users = User.objects.all()
    return render(request, 'user_list.html', {'users': users})


def word_gener(request, pk):
    form = WordGenForm(request.POST)
    user = User.objects.get(id=pk)
    print(form['data'].data)
    indexes = Index.objects.filter(teacher=user, data=form['data'].data)
    document = Document()


    for education in Education.objects.all():
        document.add_heading(education.title, 0)
        table = document.add_table(rows=1, cols=6, style="Table Grid")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Наименование показателей'
        hdr_cells[2].text = 'балл'
        hdr_cells[3].text = 'Кол-во показателей'
        hdr_cells[4].text = 'Сумма баллов'
        hdr_cells[5].text = 'Примечание'
        for category in Category.objects.all():
            i = 0

            if category.catdisciplines.filter(teacher=user):
                row_cells = table.add_row().cells
                row_cells[0].text = category.title

                for discipline in category.catdisciplines.all():
                    print(len(discipline.indexes.filter(teacher=user)))
                    if len(discipline.indexes.filter(teacher=user)) != 0:
                        if discipline in education.edudisciplines.all():
                            for index in indexes:
                                if index.discripline == discipline:
                                    i += 1
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = str(i)
                                    row_cells[1].text = index.discripline.title
                                    row_cells[2].text = str(index.bally)
                                    row_cells[3].text = str(index.quantity)
                                    row_cells[4].text = str(index.bally * index.quantity)
                                    row_cells[5].text = index.primechanie
            else:
                print("no")

    oldtime = time.time()
    docume = document.save(f'media/{user.username}-{form["data"].data}-{oldtime}.docx')
    das = DocumentOtchet(user=user, data=form["data"].data, document=f'{user.username}-{form["data"].data}-{oldtime}.docx')
    das = das.save()
    return redirect(str(f'/media/{user.username}-{form["data"].data}-{oldtime}.docx/?next=/'))